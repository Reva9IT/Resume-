from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_docx(data):
    doc = Document()

    doc.add_heading(data["name"], 0)

    #doc.add_paragraph(f"Email: {data['email']} | Phone: {data['phone']}")
    #doc.add_paragraph(f"Links: {data['links']}")

    #doc.add_heading("Target Role", level=1)
    #doc.add_paragraph(data["role"])

    doc.add_heading("Education", level=1)
    doc.add_paragraph(data["education"])

    doc.add_heading("Skills", level=1)
    for skill in data["skills"].split(","):
        doc.add_paragraph(skill.strip(), style='List Bullet')

    #doc.add_heading("Projects", level=1)
    #for project in data["projects"].split(","):
       # doc.add_paragraph(project.strip(), style='List Bullet')

    doc.add_heading("Experience", level=1)
    doc.add_paragraph(data["experience"])

    #doc.add_heading("Achievements", level=1)
    #doc.add_paragraph(data["achievements"])

    #doc.add_heading("Summary", level=1)
    #doc.add_paragraph(data["summary"])

    file_path = "resume.docx"
    doc.save(file_path)
    return file_path


def create_pdf(data):
    file_path = "resume.pdf"
    c = canvas.Canvas(file_path, pagesize=letter)

    y = 750

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, data["name"])
    y -= 20

    c.setFont("Helvetica", 10)
    c.drawString(50, y, f"Email: {data['email']} | Phone: {data['phone']}")
    y -= 15
    c.drawString(50, y, f"Links: {data['links']}")
    y -= 25

    c.setFont("Helvetica", 12)

    sections = [
        #("Role", data["role"]),
        ("Education", data["education"]),
        ("Skills", data["skills"]),
        #("Projects", data["projects"]),
        ("Experience", data["experience"]),
        #("Achievements", data["achievements"]),
        #("Summary", data["summary"]),
    ]

    for title, content in sections:
        c.drawString(50, y, title + ":")
        y -= 15

        for line in content.split(","):
            c.drawString(70, y, "- " + line.strip())
            y -= 12

        y -= 10

    c.save()
    return file_path
